"""Shared Markdown-to-document export helpers for writeAgent skills."""

from __future__ import annotations

import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET


JOS_TEMPLATE_PROFILE = "journal_of_software_2025"
JOS_TEMPLATE_SOURCE = "case/references/软件学报排版样例2025年版.doc"
JOS_BODY_FONT_SIZE_PT = 9
JOS_BODY_LINE_SPACING_PT = 12
JOS_HEADER_FONT_SIZE_PT = 7.5
JOS_TITLE_FONT_SIZE_PT = 14
JOS_TITLE_LINE_SPACING_PT = 24
JOS_AFFILIATION_FONT_SIZE_PT = 8
_CITATION_RUN_RE = re.compile(r"\[(?:\d+(?:\s*[,，\-–—−]\s*\d+)*)\]")
_REFERENCE_HEADING_RE = re.compile(r"^(?:参考文献|References)$", re.IGNORECASE)
_REFERENCE_LINE_RE = re.compile(r"^(\[\d+\]\s+)(.*)$")
_NUMBERED_REFERENCE_LINE_RE = re.compile(r"^(\d+)[.)]\s+(.*)$")
_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
_JOS_PUNCTUATION_TRANSLATION = str.maketrans(
    {
        "，": ",",
        "。": ".",
        "；": ";",
        "：": ":",
        "（": "(",
        "）": ")",
        "【": "[",
        "】": "]",
        "［": "[",
        "］": "]",
        "｛": "{",
        "｝": "}",
        "“": '"',
        "”": '"',
        "‘": "'",
        "’": "'",
        "、": ",",
        "《": "<",
        "》": ">",
        "？": "?",
        "！": "!",
        "—": "-",
        "–": "-",
        "－": "-",
        "～": "~",
        "…": "...",
        "·": ".",
    }
)


class DocumentExportError(RuntimeError):
    """Raised when a required export cannot be produced."""


def export_markdown_document(
    markdown: str,
    output_base: str | Path,
    *,
    title: str | None = None,
    template_profile: str | None = None,
    template_source_path: str | None = None,
) -> dict[str, Any]:
    """Export a Markdown paper to DOCX and, when possible, PDF.

    DOCX is a required artifact for Skill5/6, so DOCX failures raise
    ``DocumentExportError``. PDF export is optional: missing dependencies,
    fonts, or rendering issues are recorded as ``unavailable`` in the returned
    status instead of failing the skill contract.
    """

    profile = normalize_template_profile(template_profile, template_source_path)
    base = Path(output_base)
    base.parent.mkdir(parents=True, exist_ok=True)
    docx_path = base.with_suffix(".docx")
    pdf_path = base.with_suffix(".pdf")

    _write_docx(markdown, docx_path, title=title, template_profile=profile)
    export_status: dict[str, Any] = {
        "docx": {
            "status": "generated",
            "path": str(docx_path),
        }
    }

    exported_pdf_path = _export_pdf(markdown, docx_path, pdf_path, title=title, template_profile=profile)
    export_status["pdf"] = exported_pdf_path["status"]

    return {
        "docx_path": str(docx_path),
        "pdf_path": exported_pdf_path["path"],
        "export_status": export_status,
        "template_profile": profile,
        "template_source_path": template_source_path or (JOS_TEMPLATE_SOURCE if profile == JOS_TEMPLATE_PROFILE else None),
        "template_conformance_report": build_template_conformance_report(docx_path, profile),
    }


def normalize_template_profile(
    template_profile: str | None,
    template_source_path: str | None = None,
) -> str | None:
    profile = str(template_profile or "").strip()
    if profile:
        return profile
    source = str(template_source_path or "").strip()
    if "软件学报排版样例" in source or "journal_of_software" in source.lower():
        return JOS_TEMPLATE_PROFILE
    sample = Path(JOS_TEMPLATE_SOURCE)
    if sample.exists():
        return JOS_TEMPLATE_PROFILE
    return None


def build_template_conformance_report(docx_path: str | Path, template_profile: str | None) -> dict[str, Any]:
    profile = normalize_template_profile(template_profile)
    if profile != JOS_TEMPLATE_PROFILE:
        return {"profile": profile or "default", "checks": {}, "passed": True}

    checks = {
        "body_citations_superscript": False,
        "reference_numbers_not_superscript": False,
        "page_setup_matches_template": False,
        "heading_styles_match_template": False,
        "body_line_spacing_matches_template": False,
    }
    try:
        checks["body_citations_superscript"] = _body_citations_are_superscript(Path(docx_path))
        checks["reference_numbers_not_superscript"] = _reference_numbers_are_not_superscript(Path(docx_path))
        checks["page_setup_matches_template"] = _page_setup_matches_jos(Path(docx_path))
        checks["heading_styles_match_template"] = _heading_styles_match_jos(Path(docx_path))
        checks["body_line_spacing_matches_template"] = _body_line_spacing_matches_jos(Path(docx_path))
    except Exception as exc:  # pragma: no cover - report should be best effort.
        return {
            "profile": profile,
            "checks": checks,
            "passed": False,
            "error": str(exc),
        }
    return {
        "profile": profile,
        "checks": checks,
        "passed": all(checks.values()),
    }


def _write_docx(markdown: str, path: Path, *, title: str | None, template_profile: str | None) -> None:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - exercised in integration environments.
        raise DocumentExportError("python-docx is required to export DOCX") from exc

    try:
        document = Document()
        _configure_docx_styles(document, template_profile=template_profile)
        if template_profile == JOS_TEMPLATE_PROFILE:
            _add_jos_journal_header(document, template_profile=template_profile)
        first_heading_written = False
        in_references = False
        pending_role: str | None = None
        jos_front_matter_tail_added = False
        pending: list[str] = []

        def flush_pending() -> None:
            nonlocal pending_role
            if not pending:
                return
            paragraph_text = " ".join(line.strip() for line in pending if line.strip()).strip()
            pending.clear()
            if paragraph_text:
                if pending_role == "abstract":
                    _add_jos_labeled_paragraph(
                        document,
                        "摘  要:",
                        paragraph_text,
                        template_profile=template_profile,
                        east_asia_font="KaiTi",
                    )
                else:
                    _add_body_paragraph(document, paragraph_text, template_profile=template_profile)
            pending_role = None

        def ensure_jos_front_matter_tail() -> None:
            nonlocal jos_front_matter_tail_added
            if template_profile != JOS_TEMPLATE_PROFILE or jos_front_matter_tail_added:
                return
            _add_jos_publication_metadata(
                document,
                title=title or _first_heading(markdown) or "论文题名待补充",
                template_profile=template_profile,
            )
            jos_front_matter_tail_added = True

        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                flush_pending()
                continue

            heading = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
            if heading:
                flush_pending()
                level = len(heading.group(1))
                text = _strip_inline_markdown(heading.group(2))
                in_references = bool(_REFERENCE_HEADING_RE.match(text))
                if template_profile == JOS_TEMPLATE_PROFILE and text in {"摘要", "摘 要", "摘  要"}:
                    pending_role = "abstract"
                    continue
                if template_profile == JOS_TEMPLATE_PROFILE and first_heading_written:
                    ensure_jos_front_matter_tail()
                _add_heading(
                    document,
                    text or title or "",
                    level=level,
                    first_heading=level == 1 and not first_heading_written,
                    template_profile=template_profile,
                    is_references=in_references,
                )
                if level == 1 and not first_heading_written:
                    first_heading_written = True
                continue

            if template_profile == JOS_TEMPLATE_PROFILE and _is_keywords_line(line):
                flush_pending()
                keywords_text = _strip_keywords_label(_strip_inline_markdown(line))
                _add_jos_labeled_paragraph(
                    document,
                    "关键词:",
                    keywords_text,
                    template_profile=template_profile,
                )
                ensure_jos_front_matter_tail()
                continue

            if in_references or _is_reference_line(line):
                flush_pending()
                ensure_jos_front_matter_tail()
                _add_reference_paragraph(document, _strip_inline_markdown(line), template_profile=template_profile)
                continue

            if _is_list_line(line):
                flush_pending()
                ensure_jos_front_matter_tail()
                paragraph = document.add_paragraph(style="List Bullet")
                _add_runs_with_optional_superscript(
                    paragraph,
                    _strip_inline_markdown(line.lstrip("-*0123456789. ")),
                    superscript_citations=template_profile == JOS_TEMPLATE_PROFILE,
                    template_profile=template_profile,
                )
                _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile)
                continue

            if _is_table_line(line):
                flush_pending()
                ensure_jos_front_matter_tail()
                paragraph = document.add_paragraph()
                _add_runs_with_optional_superscript(
                    paragraph,
                    _strip_inline_markdown(line),
                    superscript_citations=template_profile == JOS_TEMPLATE_PROFILE,
                    template_profile=template_profile,
                )
                _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile)
                continue

            pending.append(line)

        flush_pending()
        if template_profile == JOS_TEMPLATE_PROFILE and first_heading_written:
            ensure_jos_front_matter_tail()
        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
    except DocumentExportError:
        raise
    except Exception as exc:  # pragma: no cover - depends on python-docx internals.
        raise DocumentExportError(f"failed to export DOCX: {exc}") from exc


def _configure_docx_styles(document: Any, *, template_profile: str | None) -> None:
    from docx.shared import Cm, Pt

    if template_profile == JOS_TEMPLATE_PROFILE:
        section = document.sections[0]
        section.page_width = Cm(18.4)
        section.page_height = Cm(26.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        _set_style_font(document.styles["Normal"], "Times New Roman", "SimSun", JOS_BODY_FONT_SIZE_PT)
        _set_exact_line_spacing(document.styles["Normal"].paragraph_format, JOS_BODY_LINE_SPACING_PT)
        _set_style_font(document.styles["Title"], "Times New Roman", "SimHei", JOS_TITLE_FONT_SIZE_PT)
        _set_exact_line_spacing(document.styles["Title"].paragraph_format, JOS_TITLE_LINE_SPACING_PT)
        _remove_style_paragraph_borders(document.styles["Title"])
        _set_style_font(document.styles["Heading 1"], "Times New Roman", "SimHei", 10.5)
        _set_style_font(document.styles["Heading 2"], "Times New Roman", "SimHei", JOS_BODY_FONT_SIZE_PT)
        _set_style_font(document.styles["Heading 3"], "Times New Roman", "SimHei", JOS_BODY_FONT_SIZE_PT)
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            _set_exact_line_spacing(document.styles[style_name].paragraph_format, JOS_BODY_LINE_SPACING_PT)
        return

    chinese_font = "SimSun"
    normal = document.styles["Normal"]
    normal.font.name = chinese_font
    normal.font.size = Pt(12)
    _set_east_asia_font(normal.font, chinese_font)
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = document.styles[style_name]
        except KeyError:
            continue
        style.font.name = chinese_font
        _set_east_asia_font(style.font, chinese_font)


def _set_style_font(style: Any, ascii_font: str, east_asia_font: str, size_pt: float) -> None:
    from docx.shared import Pt, RGBColor

    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(style.font, east_asia_font)


def _set_exact_line_spacing(paragraph_format: Any, spacing_pt: float) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    paragraph_format.line_spacing = Pt(spacing_pt)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _remove_style_paragraph_borders(style: Any) -> None:
    try:
        from docx.oxml.ns import qn

        p_pr = style.element.get_or_add_pPr()
        while True:
            border = p_pr.find(qn("w:pBdr"))
            if border is None:
                return
            p_pr.remove(border)
    except Exception:
        return


def _set_east_asia_font(font: Any, font_name: str) -> None:
    try:
        from docx.oxml.ns import qn

        font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        return


def _add_heading(
    document: Any,
    text: str,
    *,
    level: int,
    first_heading: bool,
    template_profile: str | None,
    is_references: bool,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if template_profile == JOS_TEMPLATE_PROFILE:
        if first_heading:
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _set_exact_line_spacing(paragraph.paragraph_format, JOS_TITLE_LINE_SPACING_PT)
            _add_plain_run(
                paragraph,
                text,
                template_profile=template_profile,
                size_pt=JOS_TITLE_FONT_SIZE_PT,
                east_asia_font="SimHei",
                bold=True,
            )
            _add_jos_author_placeholders(document, template_profile=template_profile)
            return
        if is_references:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile, heading=True)
            _add_plain_run(paragraph, text, template_profile=template_profile, size_pt=9, east_asia_font="SimHei", bold=True)
            return
        mapped_level = max(1, min(level - 1, 3))
        paragraph = document.add_heading("", mapped_level)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_plain_run(
            paragraph,
            text,
            template_profile=template_profile,
            size_pt=10.5 if mapped_level == 1 else 9,
            east_asia_font="SimHei",
        )
        return

    if first_heading:
        paragraph = document.add_heading(text, 0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_heading(text, min(level, 4))


def _add_body_paragraph(document: Any, text: str, *, template_profile: str | None) -> None:
    paragraph = document.add_paragraph()
    _add_runs_with_optional_superscript(
        paragraph,
        _strip_inline_markdown(text),
        superscript_citations=template_profile == JOS_TEMPLATE_PROFILE,
        template_profile=template_profile,
    )
    _format_paragraph(paragraph, template_profile=template_profile)


def _add_jos_journal_header(document: Any, *, template_profile: str | None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.0)
    table.columns[1].width = Cm(7.0)
    _remove_table_borders(table)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for text in (
        "软件学报 ISSN 1000-9825, CODEN RUXUEW",
        "Journal of Software, doi: 10.13328/j.cnki.jos.000000",
        "©中国科学院软件研究所版权所有.",
    ):
        if left.text:
            left.add_run().add_break()
        _add_plain_run(left, text, template_profile=template_profile, size_pt=JOS_HEADER_FONT_SIZE_PT)
    for text in ("E-mail: jos@iscas.ac.cn", "http://www.jos.org.cn", "Tel: +86-10-62562563"):
        if right.text:
            right.add_run().add_break()
        _add_plain_run(right, text, template_profile=template_profile, size_pt=JOS_HEADER_FONT_SIZE_PT)
    for paragraph in (left, right):
        paragraph.paragraph_format.space_after = Pt(0)
        _set_exact_line_spacing(paragraph.paragraph_format, JOS_BODY_LINE_SPACING_PT)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


def _remove_table_borders(table: Any) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "nil")
            borders.append(element)
        table._tbl.tblPr.append(borders)
    except Exception:
        return


def _add_jos_author_placeholders(document: Any, *, template_profile: str | None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    for text, size, east_asia_font in (
        ("作者信息待补充", 10.5, "FangSong"),
        ("(单位信息待补充)", JOS_AFFILIATION_FONT_SIZE_PT, "SimSun"),
        ("通讯作者: 待补充, E-mail: 待补充", 9, "SimSun"),
    ):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if text != "通讯作者: 待补充, E-mail: 待补充" else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(4)
        _set_exact_line_spacing(paragraph.paragraph_format, JOS_BODY_LINE_SPACING_PT)
        _add_plain_run(
            paragraph,
            text,
            template_profile=template_profile,
            size_pt=size,
            east_asia_font=east_asia_font,
        )


def _add_jos_labeled_paragraph(
    document: Any,
    label: str,
    text: str,
    *,
    template_profile: str | None,
    east_asia_font: str | None = None,
) -> None:
    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile)
    _add_plain_run(
        paragraph,
        f"{label}  ",
        template_profile=template_profile,
        bold=True,
        east_asia_font=east_asia_font,
    )
    _add_runs_with_optional_superscript(
        paragraph,
        text,
        superscript_citations=True,
        template_profile=template_profile,
        east_asia_font=east_asia_font,
    )


def _add_jos_publication_metadata(
    document: Any,
    *,
    title: str,
    template_profile: str | None,
) -> None:
    _add_jos_labeled_paragraph(document, "中图法分类号:", "TP311", template_profile=template_profile)
    _add_jos_citation_paragraph(
        document,
        f"中文引用格式: 作者信息待补充. {title}[J]. 软件学报, 2025, 0(0): 1-15. http://www.jos.org.cn/1000-9825/0000.htm",
        template_profile=template_profile,
    )
    _add_jos_citation_paragraph(
        document,
        "英文引用格式: Author information to be supplied. English title to be supplied[J]. "
        "Ruan Jian Xue Bao/Journal of Software, 2025, 0(0): 1-15. http://www.jos.org.cn/1000-9825/0000.htm",
        template_profile=template_profile,
    )
    _add_jos_english_title(document, "English Title To Be Supplied", template_profile=template_profile)
    _add_jos_english_authors(document, template_profile=template_profile)
    _add_jos_labeled_paragraph(
        document,
        "Abstract:",
        "English abstract to be supplied.",
        template_profile=template_profile,
    )
    _add_jos_labeled_paragraph(
        document,
        "Key words:",
        "English key words to be supplied.",
        template_profile=template_profile,
    )


def _add_jos_citation_paragraph(document: Any, text: str, *, template_profile: str | None) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    _shade_paragraph(paragraph, "EDEDED")
    _add_plain_run(paragraph, text, template_profile=template_profile, size_pt=8)


def _add_jos_english_title(document: Any, text: str, *, template_profile: str | None) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    _set_exact_line_spacing(paragraph.paragraph_format, JOS_BODY_LINE_SPACING_PT)
    _add_plain_run(paragraph, text, template_profile=template_profile, size_pt=14, bold=True)


def _add_jos_english_authors(document: Any, *, template_profile: str | None) -> None:
    for text in (
        "AUTHOR INFORMATION TO BE SUPPLIED",
        "(Affiliation information to be supplied)",
    ):
        paragraph = document.add_paragraph()
        _format_paragraph(paragraph, first_line_indent=False, template_profile=template_profile)
        _add_plain_run(paragraph, text, template_profile=template_profile, size_pt=9)


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_pr = paragraph._p.get_or_add_pPr()
        existing = p_pr.find(qn("w:shd"))
        if existing is not None:
            p_pr.remove(existing)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        p_pr.append(shading)
    except Exception:
        return


def _add_reference_paragraph(document: Any, text: str, *, template_profile: str | None) -> None:
    paragraph = document.add_paragraph()
    match = _REFERENCE_LINE_RE.match(text)
    if match:
        _add_plain_run(paragraph, match.group(1), template_profile=template_profile, size_pt=7.5 if template_profile == JOS_TEMPLATE_PROFILE else None)
        _add_plain_run(paragraph, match.group(2), template_profile=template_profile, size_pt=7.5 if template_profile == JOS_TEMPLATE_PROFILE else None)
    else:
        numbered_match = _NUMBERED_REFERENCE_LINE_RE.match(text)
        if numbered_match and template_profile == JOS_TEMPLATE_PROFILE:
            _add_plain_run(paragraph, f"[{numbered_match.group(1)}] ", template_profile=template_profile, size_pt=7.5)
            _add_plain_run(paragraph, numbered_match.group(2), template_profile=template_profile, size_pt=7.5)
        else:
            _add_plain_run(paragraph, text, template_profile=template_profile, size_pt=7.5 if template_profile == JOS_TEMPLATE_PROFILE else None)
    _format_paragraph(
        paragraph,
        first_line_indent=False,
        hanging_reference=True,
        template_profile=template_profile,
    )


def _add_runs_with_optional_superscript(
    paragraph: Any,
    text: str,
    *,
    superscript_citations: bool,
    template_profile: str | None,
    east_asia_font: str | None = None,
) -> None:
    position = 0
    for match in _CITATION_RUN_RE.finditer(text):
        if match.start() > position:
            _add_plain_run(
                paragraph,
                text[position : match.start()],
                template_profile=template_profile,
                east_asia_font=east_asia_font,
            )
        run = _add_plain_run(
            paragraph,
            match.group(0),
            template_profile=template_profile,
            east_asia_font=east_asia_font,
        )
        if superscript_citations:
            run.font.superscript = True
        position = match.end()
    if position < len(text):
        _add_plain_run(
            paragraph,
            text[position:],
            template_profile=template_profile,
            east_asia_font=east_asia_font,
        )


def _add_plain_run(
    paragraph: Any,
    text: str,
    *,
    template_profile: str | None,
    size_pt: float | None = None,
    east_asia_font: str | None = None,
    bold: bool = False,
) -> Any:
    from docx.shared import Pt

    run = paragraph.add_run(_normalize_jos_text(text) if template_profile == JOS_TEMPLATE_PROFILE else text)
    if template_profile == JOS_TEMPLATE_PROFILE:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt or 9)
        run.font.bold = bold
        _set_east_asia_font(run.font, east_asia_font or "SimSun")
    return run


def _normalize_jos_text(text: str) -> str:
    return text.translate(_JOS_PUNCTUATION_TRANSLATION)


def _format_paragraph(
    paragraph: Any,
    *,
    first_line_indent: bool = True,
    hanging_reference: bool = False,
    template_profile: str | None,
    heading: bool = False,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    if template_profile == JOS_TEMPLATE_PROFILE:
        _set_exact_line_spacing(paragraph.paragraph_format, JOS_BODY_LINE_SPACING_PT)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if heading else WD_ALIGN_PARAGRAPH.JUSTIFY
        if hanging_reference:
            paragraph.paragraph_format.left_indent = Pt(20.1)
            paragraph.paragraph_format.first_line_indent = Pt(-20.1)
        elif first_line_indent:
            paragraph.paragraph_format.first_line_indent = Pt(15.6)
        return

    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Cm(0.08)
    if hanging_reference:
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
    elif first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)


def _export_pdf(
    markdown: str,
    docx_path: Path,
    pdf_path: Path,
    *,
    title: str | None,
    template_profile: str | None,
) -> dict[str, Any]:
    method = "reportlab"
    fallback_reason: str | None = None
    try:
        if template_profile == JOS_TEMPLATE_PROFILE:
            _write_pdf_from_docx_with_word(docx_path, pdf_path)
            method = "word_com_docx"
        else:
            _write_pdf(markdown, pdf_path, title=title, template_profile=template_profile)
    except Exception as first_exc:
        try:
            _write_pdf(markdown, pdf_path, title=title, template_profile=template_profile)
            method = "reportlab_fallback"
            fallback_reason = str(first_exc)
        except Exception as fallback_exc:  # PDF is explicitly best-effort.
            return {
                "status": {
                    "status": "unavailable",
                    "path": None,
                    "reason": f"{first_exc}; fallback failed: {fallback_exc}",
                },
                "path": None,
            }
    return {
        "status": {
            "status": "generated",
            "path": str(pdf_path),
            "method": method,
            **({"fallback_reason": fallback_reason} if fallback_reason else {}),
        },
        "path": str(pdf_path),
    }


def _write_pdf_from_docx_with_word(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="writeagent-word-pdf-") as temp_dir:
        safe_docx_path = Path(temp_dir) / "source.docx"
        safe_pdf_path = Path(temp_dir) / "output.pdf"
        shutil.copyfile(docx_path, safe_docx_path)
        _write_pdf_from_safe_docx_with_word(safe_docx_path, safe_pdf_path)
        shutil.copyfile(safe_pdf_path, pdf_path)


def _write_pdf_from_safe_docx_with_word(docx_path: Path, pdf_path: Path) -> None:
    script = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $null
try {{
  $doc = $word.Documents.Open({json.dumps(str(docx_path.resolve()))}, $false, $true)
  $doc.SaveAs([ref]{json.dumps(str(pdf_path.resolve()))}, [ref]17)
}} finally {{
  if ($doc -ne $null) {{ $doc.Close($false) }}
  $word.Quit()
}}
"""
    result = subprocess.run(
        ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError((result.stderr or result.stdout or "Word COM PDF conversion failed").strip())


def _write_pdf(markdown: str, path: Path, *, title: str | None, template_profile: str | None = None) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfgen import canvas
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed") from exc

    font_name = _register_pdf_font(pdfmetrics)
    if template_profile == JOS_TEMPLATE_PROFILE:
        page_width, page_height = 18.4 * cm, 26.0 * cm
        left = 1.5 * cm
        right = page_width - 1.5 * cm
        top = page_height - 1.0 * cm
        bottom = 2.2 * cm
        body_font_size = JOS_BODY_FONT_SIZE_PT
        body_line_height = JOS_BODY_LINE_SPACING_PT
    else:
        page_width, page_height = A4
        left = 54
        right = page_width - 54
        top = page_height - 54
        bottom = 54
        body_font_size = 10.5
        body_line_height = 16.5
    y = top
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=(page_width, page_height))
    c.setTitle(title or _first_heading(markdown) or "writeAgent paper")

    def new_page() -> None:
        nonlocal y
        c.showPage()
        c.setPageSize((page_width, page_height))
        c.setFont(font_name, body_font_size)
        y = top

    def draw_wrapped(text: str, *, font_size: float, line_height: float, align_center: bool = False) -> None:
        nonlocal y
        c.setFont(font_name, font_size)
        for piece in _wrap_pdf_text(text, c, font_name=font_name, font_size=font_size, max_width=right - left):
            if y - line_height < bottom:
                new_page()
                c.setFont(font_name, font_size)
            if align_center:
                c.drawCentredString(page_width / 2, y, piece)
            else:
                c.drawString(left, y, piece)
            y -= line_height

    c.setFont(font_name, body_font_size)
    first_heading = True
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            y -= body_line_height * 0.55
            if y < bottom:
                new_page()
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        font_size = body_font_size
        line_height = body_line_height
        align_center = False
        if heading:
            line = heading.group(2)
            level = len(heading.group(1))
            if template_profile == JOS_TEMPLATE_PROFILE:
                font_size = JOS_TITLE_FONT_SIZE_PT if first_heading and level == 1 else (10.5 if level <= 2 else JOS_BODY_FONT_SIZE_PT)
                line_height = 24 if first_heading and level == 1 else JOS_BODY_LINE_SPACING_PT
                align_center = False
            else:
                font_size = max(12, 18 - level)
                line_height = font_size + 6
            first_heading = False
        output_line = _strip_inline_markdown(line)
        if template_profile == JOS_TEMPLATE_PROFILE:
            output_line = _normalize_jos_text(output_line)
        draw_wrapped(output_line, font_size=font_size, line_height=line_height, align_center=align_center)
        y -= body_line_height * 0.25

    c.save()


def _register_pdf_font(pdfmetrics: Any) -> str:
    try:
        from reportlab.pdfbase.ttfonts import TTFont

        for candidate in (
            r"C:\Windows\Fonts\STSONG.TTF",
            r"C:\Windows\Fonts\SourceHanSansSC-Medium.ttf",
            r"C:\Windows\Fonts\NotoSansSC-VF.ttf",
            r"C:\Windows\Fonts\simsunb.ttf",
            r"C:\Windows\Fonts\simhei.ttf",
            r"C:\Windows\Fonts\Deng.ttf",
        ):
            if Path(candidate).exists():
                font_name = "WriteAgentCJK"
                pdfmetrics.registerFont(TTFont(font_name, candidate))
                return font_name
    except Exception:
        pass

    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def _body_citations_are_superscript(docx_path: Path) -> bool:
    paragraphs = _read_docx_paragraph_runs(docx_path)
    seen = False
    in_references = False
    for runs in paragraphs:
        paragraph_text = "".join(text for text, _ in runs)
        if _REFERENCE_HEADING_RE.match(paragraph_text.strip()):
            in_references = True
            continue
        if in_references:
            continue
        for text, superscript in runs:
            if _CITATION_RUN_RE.fullmatch(text):
                seen = True
                if not superscript:
                    return False
    return seen


def _reference_numbers_are_not_superscript(docx_path: Path) -> bool:
    paragraphs = _read_docx_paragraph_runs(docx_path)
    in_references = False
    seen = False
    for runs in paragraphs:
        paragraph_text = "".join(text for text, _ in runs)
        if _REFERENCE_HEADING_RE.match(paragraph_text.strip()):
            in_references = True
            continue
        if not in_references:
            continue
        for text, superscript in runs:
            if re.match(r"^\[\d+\]", text):
                seen = True
                if superscript:
                    return False
    return seen


def _page_setup_matches_jos(docx_path: Path) -> bool:
    from docx import Document

    section = Document(str(docx_path)).sections[0]
    return (
        round(section.page_width.cm, 1) == 18.4
        and round(section.page_height.cm, 1) == 26.0
        and round(section.top_margin.cm, 1) == 1.0
        and round(section.bottom_margin.cm, 1) == 2.2
        and round(section.left_margin.cm, 1) == 1.5
        and round(section.right_margin.cm, 1) == 1.5
    )


def _heading_styles_match_jos(docx_path: Path) -> bool:
    from docx import Document

    document = Document(str(docx_path))
    return (
        document.styles["Normal"].font.size.pt == JOS_BODY_FONT_SIZE_PT
        and document.styles["Heading 1"].font.size.pt == 10.5
        and document.styles["Heading 2"].font.size.pt == JOS_BODY_FONT_SIZE_PT
    )


def _body_line_spacing_matches_jos(docx_path: Path) -> bool:
    checked = False
    with zipfile.ZipFile(docx_path) as archive:
        document_root = ET.fromstring(archive.read("word/document.xml"))
        styles_root = ET.fromstring(archive.read("word/styles.xml"))

    for paragraph in document_root.findall(".//w:p", _NS):
        if not "".join(node.text or "" for node in paragraph.findall(".//w:t", _NS)).strip():
            continue
        spacing = paragraph.find("w:pPr/w:spacing", _NS)
        if spacing is None:
            continue
        checked = True
        style = paragraph.find("w:pPr/w:pStyle", _NS)
        style_id = style.attrib.get(f"{{{_NS['w']}}}val") if style is not None else None
        if not _line_spacing_is_expected(spacing, _expected_line_spacing_twips(style_id)):
            return False

    for style_id in ("Normal", "Title", "Heading1", "Heading2", "Heading3"):
        style = styles_root.find(f".//w:style[@w:styleId='{style_id}']", _NS)
        if style is None:
            continue
        spacing = style.find("w:pPr/w:spacing", _NS)
        if spacing is None:
            continue
        checked = True
        if not _line_spacing_is_expected(spacing, _expected_line_spacing_twips(style_id)):
            return False

    return checked


def _expected_line_spacing_twips(style_id: str | None) -> str:
    spacing_pt = JOS_TITLE_LINE_SPACING_PT if style_id == "Title" else JOS_BODY_LINE_SPACING_PT
    return str(int(spacing_pt * 20))


def _line_spacing_is_expected(spacing: ET.Element, expected_twips: str) -> bool:
    line = spacing.attrib.get(f"{{{_NS['w']}}}line")
    line_rule = spacing.attrib.get(f"{{{_NS['w']}}}lineRule")
    return line == expected_twips and line_rule == "exact"


def _read_docx_paragraph_runs(docx_path: Path) -> list[list[tuple[str, bool]]]:
    with zipfile.ZipFile(docx_path) as archive:
        root = ET.fromstring(archive.read("word/document.xml"))
    paragraphs: list[list[tuple[str, bool]]] = []
    for paragraph in root.findall(".//w:p", _NS):
        runs: list[tuple[str, bool]] = []
        for run in paragraph.findall("w:r", _NS):
            text = "".join(node.text or "" for node in run.findall("w:t", _NS))
            if not text:
                continue
            vert = run.find("w:rPr/w:vertAlign", _NS)
            superscript = vert is not None and vert.attrib.get(f"{{{_NS['w']}}}val") == "superscript"
            runs.append((text, superscript))
        if runs:
            paragraphs.append(runs)
    return paragraphs


def _strip_inline_markdown(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def _is_keywords_line(line: str) -> bool:
    stripped = _strip_inline_markdown(line).strip()
    return bool(re.match(r"^(?:关键词|关键字)\s*[:：]", stripped))


def _strip_keywords_label(text: str) -> str:
    return re.sub(r"^(?:关键词|关键字)\s*[:：]\s*", "", text).strip()


def _is_list_line(line: str) -> bool:
    return bool(re.match(r"\s*(?:[-*]\s+|\d+[.)]\s+)", line))


def _is_reference_line(line: str) -> bool:
    return bool(_REFERENCE_LINE_RE.match(line.strip()))


def _is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _first_heading(markdown: str) -> str | None:
    for line in markdown.splitlines():
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return _strip_inline_markdown(match.group(1))
    return None


def _wrap_pdf_text(text: str, canvas_obj: Any, *, font_name: str, font_size: float, max_width: float) -> list[str]:
    if not text:
        return [""]

    def width(value: str) -> float:
        return canvas_obj.stringWidth(value, font_name, font_size)

    lines: list[str] = []
    current = ""
    for token in _pdf_wrap_tokens(text):
        candidate = f"{current}{token}"
        if not current or width(candidate) <= max_width:
            current = candidate
            continue
        lines.append(current.rstrip())
        current = token.lstrip()
        while current and width(current) > max_width:
            split_at = _largest_prefix_that_fits(current, width, max_width)
            lines.append(current[:split_at].rstrip())
            current = current[split_at:].lstrip()
    if current:
        lines.append(current.rstrip())
    return lines or [text]


def _pdf_wrap_tokens(text: str) -> list[str]:
    return re.findall(r"[A-Za-z0-9]+(?:[-_/.:][A-Za-z0-9]+)*|[\u4e00-\u9fff]|.", text)


def _largest_prefix_that_fits(text: str, width: Any, max_width: float) -> int:
    for index in range(1, len(text) + 1):
        if width(text[:index]) > max_width:
            return max(1, index - 1)
    return len(text)
